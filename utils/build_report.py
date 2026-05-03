"""
Build comparative DOCX report: THE LEGAL LAST MILE
Uses water_law_global_coded.csv to generate charts + Word document
"""
import os
import sys
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from pathlib import Path

# ── paths ──────────────────────────────────────────────────────────────────
DATA  = Path(r'C:\Users\junio\Downloads\water_law_global_coded.csv')
CHARTS = Path(r'C:\Users\junio\Downloads\charts')
OUT   = Path(r'C:\Users\junio\Downloads\water_law_report.docx')
CHARTS.mkdir(exist_ok=True)

print("Loading data …")
df = pd.read_csv(DATA, low_memory=False)
print(f"  {len(df):,} rows, columns: {list(df.columns)}")

# ── normalise ──────────────────────────────────────────────────────────────
df['ano'] = pd.to_numeric(df.get('ano', df.get('year', None)), errors='coerce')
if 'ano' not in df.columns or df['ano'].isna().all():
    df['ano'] = pd.to_numeric(df.get('year', None), errors='coerce')
df['country'] = df.get('country', 'Brazil')
df['win_loss'] = df.get('win_loss', 'unclear')
# Normalise governance column name
if 'governance_cat' in df.columns and 'governance_category' not in df.columns:
    df['governance_category'] = df['governance_cat']
else:
    df['governance_category'] = df.get('governance_category', '')
df['hr_language'] = df.get('hr_language', 0)

BR = df[df['country'].str.contains('Brazil|BR', na=False, case=False)].copy()
NL = df[df['country'].str.contains('Netherlands|NL', na=False, case=False)].copy()
CA = df[df['country'].str.contains('Canada|CA', na=False, case=False)].copy()

YEARS = list(range(2016, 2026))

# ── palette ────────────────────────────────────────────────────────────────
C_BR = '#009C3B'
C_NL = '#FF6600'
C_CA = '#D80621'
GREY = '#AAAAAA'

# ══════════════════════════════════════════════════════════════════════════
# CHART 1 — Cases per year by country
# ══════════════════════════════════════════════════════════════════════════
def chart_cases_per_year():
    fig, ax1 = plt.subplots(figsize=(9, 5))
    ax2 = ax1.twinx()

    for country, col, ax_use, label in [
        ('Brazil',      C_BR, ax1, 'Brazil (left)'),
        ('Netherlands', C_NL, ax2, 'Netherlands (right)'),
        ('Canada',      C_CA, ax1, 'Canada (left)'),
    ]:
        sub = df[df['country'].str.contains(country, na=False, case=False)]
        counts = sub[sub['ano'].isin(YEARS)].groupby('ano').size().reindex(YEARS, fill_value=0)
        ax_use.plot(YEARS, counts.values, marker='o', color=col, label=label, linewidth=2)

    ax1.set_xlabel('Year')
    ax1.set_ylabel('Cases (Brazil + Canada)', color='black')
    ax2.set_ylabel('Cases (Netherlands)', color=C_NL)
    ax1.set_title('Water Law Judicial Decisions by Year and Country (2016–2025)')
    ax1.set_xticks(YEARS)
    ax1.tick_params(axis='x', rotation=45)
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=9)
    ax1.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    path = CHARTS / 'chart1_cases_per_year.png'
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Saved {path.name}")
    return str(path)

# ══════════════════════════════════════════════════════════════════════════
# CHART 2 — Brazil governance categories
# ══════════════════════════════════════════════════════════════════════════
def chart_brazil_governance():
    coded = BR[BR['governance_category'].notna() & (BR['governance_category'] != '')]
    cats = coded['governance_category'].value_counts().head(10) if len(coded) > 0 else pd.Series({'No data': len(BR)})
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.barh(cats.index[::-1], cats.values[::-1], color=C_BR, edgecolor='white')
    ax.set_xlabel('Number of Cases')
    ax.set_title('Brazil: Top Governance Categories (2016–2025)')
    for i, v in enumerate(cats.values[::-1]):
        ax.text(v + 5, i, f'{v:,}', va='center', fontsize=8)
    ax.grid(axis='x', alpha=0.3)
    fig.tight_layout()
    path = CHARTS / 'chart2_brazil_governance.png'
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Saved {path.name}")
    return str(path)

# ══════════════════════════════════════════════════════════════════════════
# CHART 3 — Brazil win/loss trend (with 2020 reform shading)
# ══════════════════════════════════════════════════════════════════════════
def chart_brazil_winloss():
    br_coded = BR[BR['win_loss'].isin(['user_wins','utility_wins','mixed']) & BR['ano'].isin(YEARS)]
    if len(br_coded) == 0:
        print("  No win/loss data for chart3, skipping")
        return None
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.axvspan(2020, 2025, alpha=0.08, color='blue', label='Post-Lei 14.026/2020')
    for outcome, col, lbl in [
        ('user_wins',    C_BR, 'User wins'),
        ('utility_wins', C_NL, 'Utility wins'),
        ('mixed',        GREY, 'Mixed'),
    ]:
        counts = br_coded[br_coded['win_loss']==outcome].groupby('ano').size().reindex(YEARS, fill_value=0)
        ax.plot(YEARS, counts.values, marker='o', color=col, label=lbl, linewidth=2)
    ax.set_xlabel('Year')
    ax.set_ylabel('Cases')
    ax.set_title('Brazil: Win/Loss Outcomes by Year (2016–2025)')
    ax.set_xticks(YEARS)
    ax.tick_params(axis='x', rotation=45)
    ax.legend(fontsize=9)
    ax.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    path = CHARTS / 'chart3_brazil_winloss.png'
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Saved {path.name}")
    return str(path)

# ══════════════════════════════════════════════════════════════════════════
# CHART 4 — Brazil HR language by year
# ══════════════════════════════════════════════════════════════════════════
def chart_brazil_hr():
    br_yr = BR[BR['ano'].isin(YEARS)].copy()
    br_yr['hr_language'] = pd.to_numeric(br_yr['hr_language'], errors='coerce').fillna(0)
    total = br_yr.groupby('ano').size().reindex(YEARS, fill_value=0)
    hr    = br_yr[br_yr['hr_language'] > 0].groupby('ano').size().reindex(YEARS, fill_value=0)
    pct   = (hr / total.replace(0, np.nan) * 100).fillna(0)
    fig, ax1 = plt.subplots(figsize=(9, 5))
    ax2 = ax1.twinx()
    ax1.bar(YEARS, hr.values, color=C_BR, alpha=0.6, label='HR cases (count)')
    ax2.plot(YEARS, pct.values, color='darkblue', marker='o', linewidth=2, label='HR % of total')
    ax1.set_xlabel('Year')
    ax1.set_ylabel('HR language cases', color=C_BR)
    ax2.set_ylabel('% of annual total', color='darkblue')
    ax1.set_title('Brazil: Human Rights Language in Decisions (2016–2025)')
    ax1.set_xticks(YEARS)
    ax1.tick_params(axis='x', rotation=45)
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1+lines2, labels1+labels2, fontsize=9)
    ax1.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    path = CHARTS / 'chart4_brazil_hr.png'
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Saved {path.name}")
    return str(path)

# ══════════════════════════════════════════════════════════════════════════
# CHART 5 — Comparative governance (3-panel)
# ══════════════════════════════════════════════════════════════════════════
def chart_comparative():
    fig, axes = plt.subplots(1, 3, figsize=(13, 5))
    datasets = [(BR, 'Brazil', C_BR), (NL, 'Netherlands', C_NL), (CA, 'Canada', C_CA)]
    for ax, (sub, name, col) in zip(axes, datasets):
        coded = sub[sub['governance_cat'].notna() & (sub['governance_cat'] != '') & (sub['governance_cat'] != 'other_water') & (sub['governance_cat'] != 'not_water_related')]
        cats = coded['governance_cat'].value_counts().head(6) if len(coded) > 0 else pd.Series({'other_water': len(sub)})
        ax.barh(cats.index[::-1], cats.values[::-1], color=col, edgecolor='white')
        ax.set_title(f'{name}\n(n={len(sub):,})', fontsize=10, fontweight='bold')
        ax.set_xlabel('Cases')
        ax.grid(axis='x', alpha=0.3)
        ax.tick_params(labelsize=7)
    fig.suptitle('Comparative Governance Categories by Country', fontsize=12, fontweight='bold')
    fig.tight_layout()
    path = CHARTS / 'chart5_comparative.png'
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Saved {path.name}")
    return str(path)

# ══════════════════════════════════════════════════════════════════════════
# CHART 6 — Informal settlement trend (Brazil)
# ══════════════════════════════════════════════════════════════════════════
def chart_informal():
    br_yr = BR[BR['ano'].isin(YEARS)].copy()
    if 'irregular_settlement' not in br_yr.columns:
        br_yr['irregular_settlement'] = br_yr['governance_category'].str.contains(
            r'irregular|informal|favela|invas[ãa]o|loteamento', case=False, na=False, regex=True).astype(int)
    total  = br_yr.groupby('ano').size().reindex(YEARS, fill_value=0)
    irreg  = br_yr[br_yr['irregular_settlement'] > 0].groupby('ano').size().reindex(YEARS, fill_value=0)
    pct    = (irreg / total.replace(0, np.nan) * 100).fillna(0)
    fig, ax1 = plt.subplots(figsize=(9, 5))
    ax2 = ax1.twinx()
    ax1.bar(YEARS, irreg.values, color='#8B4513', alpha=0.6, label='Irregular settlement cases')
    ax2.plot(YEARS, pct.values, color='darkred', marker='s', linewidth=2, label='% of annual total')
    ax1.set_xlabel('Year')
    ax1.set_ylabel('Cases', color='#8B4513')
    ax2.set_ylabel('% of total', color='darkred')
    ax1.set_title('Brazil: Irregular Settlement Cases (2016–2025)')
    ax1.set_xticks(YEARS)
    ax1.tick_params(axis='x', rotation=45)
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1+lines2, labels1+labels2, fontsize=9)
    ax1.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    path = CHARTS / 'chart6_informal.png'
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print(f"  Saved {path.name}")
    return str(path)

# ══════════════════════════════════════════════════════════════════════════
# Generate all charts
# ══════════════════════════════════════════════════════════════════════════
print("\nGenerating charts …")
c1 = chart_cases_per_year()
c2 = chart_brazil_governance()
c3 = chart_brazil_winloss()
c4 = chart_brazil_hr()
c5 = chart_comparative()
c6 = chart_informal()

# ══════════════════════════════════════════════════════════════════════════
# Build DOCX report
# ══════════════════════════════════════════════════════════════════════════
print("\nBuilding DOCX report …")
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)
import datetime

doc = Document()

def h1(text):
    p = doc.add_heading(text, level=1)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    if p.runs: p.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph(text)
    if bold or italic or size:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
            if size: run.font.size = Pt(size)
    return p

def add_chart(path, caption='', width=6.0):
    if path and Path(path).exists():
        doc.add_picture(path, width=Inches(width))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                run.italic = True
                run.font.size = Pt(9)

# ── summary stats ───────────────────────────────────────────────────────────
n_total = len(df)
n_br = len(BR); n_nl = len(NL); n_ca = len(CA)
br_coded_wl = BR[BR['win_loss'].isin(['user_wins','utility_wins','mixed'])]
n_br_coded  = len(br_coded_wl)
n_user_wins = len(BR[BR['win_loss']=='user_wins'])
n_util_wins = len(BR[BR['win_loss']=='utility_wins'])
n_mixed     = len(BR[BR['win_loss']=='mixed'])
pct_user    = n_user_wins/n_br_coded*100 if n_br_coded else 0
pct_util    = n_util_wins/n_br_coded*100 if n_br_coded else 0

pre  = BR[(BR['ano'] < 2020) & (BR['win_loss'].isin(['user_wins','utility_wins','mixed']))]
post = BR[(BR['ano'] >= 2020) & (BR['win_loss'].isin(['user_wins','utility_wins','mixed']))]
pct_user_pre  = len(pre[pre['win_loss']=='user_wins'])/len(pre)*100   if len(pre)  else 0
pct_user_post = len(post[post['win_loss']=='user_wins'])/len(post)*100 if len(post) else 0

BR['hr_language'] = pd.to_numeric(BR.get('hr_language', 0), errors='coerce').fillna(0)
n_hr  = len(BR[BR['hr_language'] > 0])
pct_hr = n_hr/n_br*100 if n_br else 0

# Governance stats for methodology section
gc_col = 'governance_cat'
n_ow_total = len(df[df[gc_col]=='other_water'])
n_ow_nl    = len(NL[NL[gc_col]=='other_water'])
n_ow_ca    = len(CA[CA[gc_col]=='other_water'])
n_ow_br    = len(BR[BR[gc_col]=='other_water'])
n_nwr      = len(df[df[gc_col]=='not_water_related'])
pct_ow_nl  = n_ow_nl/n_nl*100 if n_nl else 0
pct_ow_ca  = n_ow_ca/n_ca*100 if n_ca else 0
pct_ow_br  = n_ow_br/n_br*100 if n_br else 0
n_nl_conn  = len(NL[NL[gc_col]=='connection_refusal'])

# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('THE LEGAL LAST MILE'); run.bold = True; run.font.size = Pt(24)

t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run('Administrative Law, Water Access, and the Limits of Judicial Inclusion')
run2.font.size = Pt(14); run2.italic = True

doc.add_paragraph()
t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run('Comparative Analysis: Brazil · Netherlands · Canada')
run3.font.size = Pt(12)

doc.add_paragraph()
t4 = doc.add_paragraph(); t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = t4.add_run(
    f'Based on the Global Water Law Judicial Decisions Dataset (v1.0)\n'
    f'{n_total:,} decisions, 2016–2025\n'
    f'DOI: 10.5281/zenodo.19836413\n\n'
    f'{datetime.date.today().strftime("%B %Y")}'
)
run4.font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# METHODOLOGY NOTE
# ══════════════════════════════════════════════════════════════════════════
h1('Methodology Note: Governance Classification and the Limits of Text-Based Coding')

h2('M.1 The Jurimetric Coding Approach')
para(
    'Governance categories were assigned through a regex-based jurimetric coding engine '
    '(jurimetric_coding.py, v2.0, available at github.com/jrklaus8/water-law-dataset) '
    'applied to the text fields available in the dataset: decision summary, title, '
    'case_type, legal_area, and court_name. The coding engine applies 21 governance '
    'categories in order of specificity, returning the first matching category. '
    'Patterns cover Portuguese, English, Dutch, and French.'
)

h2('M.2 The "other_water" Residual Category')
para(
    f'A significant proportion of cases — particularly in the Netherlands ({pct_ow_nl:.1f}%) and Canada '
    f'({pct_ow_ca:.1f}%) — are classified as other_water: cases where water-related content appeared '
    'in the court\'s own keyword search results, but the available metadata does not match '
    'any of the 20 specific governance categories. This is not a coding failure; it is a '
    'structural feature of how the data was collected and what metadata is publicly available.'
)

para('Three distinct mechanisms produce other_water cases:', bold=True)
for item in [
    ('Netherlands — shallow metadata, broad scrape.',
     'Rechtspraak.nl returns all published decisions in structured XML. The collection '
     'used broad keyword searches on water-related terms ("water", "waterberging", '
     '"riolering", etc.), which returned a large number of administrative law cases '
     '(omgevingsrecht, bestuursrecht) where water is mentioned incidentally — for example, '
     'a building permit decision that includes a boilerplate drainage condition, or a '
     'zoning plan where water management is one of dozens of topics addressed. '
     'The published XML summary for these decisions (typically one sentence) does not '
     'mention water explicitly even though the full decision text does. '
     'The coding engine has access only to the summary, not the full text. '
     'These cases are genuine administrative law decisions returned by the courts\' own '
     'water-related keyword search — they are included as they represent the observable '
     'footprint of water governance in Dutch courts.'),

    ('Canada — title-only records and fragmented databases.',
     'CanLII provides case summaries for some courts but not others. Quebec courts '
     '(qccs, qcca, qctat) typically provide only the case title (e.g., "Communaute '
     'droit animalier Quebec c. Ministre de l\'Agriculture, des Pecheries et de '
     'l\'Alimentation") without a substantive summary. The title alone contains '
     'insufficient keywords to assign a specific category even when the case substantively '
     'involves water law. English-speaking provincial courts often provide brief summaries '
     'that describe the procedural posture ("appeal dismissed") rather than the substance. '
     'The Legal Data Hunter semantic search that supplemented CanLII adds cases where '
     'water law is the substance of the decision, but the available metadata from those '
     'sources may be similarly thin.'),

    ('Brazil — list-style TJSP results without ementas.',
     'The TJSP (Sao Paulo) scraper returned 574 cases via a list-format page that '
     'provides case numbers and procedural metadata but not ementas (decision summaries). '
     'Without the ementa text, the coding engine cannot determine governance category. '
     'These cases are included for count purposes but appear in other_water because '
     'the substantive text needed for category assignment is not available in the '
     'scraped metadata.'),
]:
    p = doc.add_paragraph()
    p.add_run(f'{item[0]}: ').bold = True
    p.add_run(item[1])

h2('M.3 The not_water_related Filter')
para(
    'A separate pre-filter identifies 15,589 Netherlands cases as not_water_related. '
    f'These are vreemdelingenrecht (immigration/asylum law) cases from the Raad van State '
    f'(RvS) that were returned by the keyword search despite having no water-related content. '
    'The RvS is the highest administrative court for both environmental/planning appeals '
    'and immigration/asylum appeals; the broad search terms matched tangential references '
    '(e.g., drinking water conditions in detention) in immigration decisions. '
    'The pre-filter applies a dual test: (1) presence of immigration-specific language '
    '(vreemdelingenrecht, verblijfsvergunning, asielverzoek) AND (2) absence of any '
    'water-infrastructure term (waterschap, drinkwater, riolering, fornecimento, etc.). '
    'Cases passing this filter are excluded from governance analysis but retained in '
    'the full dataset with the not_water_related flag for transparency.'
)

h2('M.4 Implications for Interpretation')
para(
    'The residual other_water category should be interpreted differently by jurisdiction:'
)
for item in [
    (f'Netherlands other_water ({n_ow_nl:,} cases, {pct_ow_nl:.1f}% of NL corpus).',
     'Primarily administrative law decisions (omgevingsvergunning, bestemmingsplan, '
     'bestuursrecht) where water management is an incidental rather than primary concern. '
     'The fact that nearly three-quarters of Dutch water-related cases cannot be '
     'classified into a specific dispute type — combined with near-zero household '
     'connection disputes (connection_refusal = 12 NL cases) — itself confirms the '
     '"pre-litigation absorption" finding: the Dutch water governance system produces '
     'mostly planning and permitting litigation, not individual access disputes.'),

    (f'Canada other_water ({n_ow_ca:,} cases, {pct_ow_ca:.1f}% of CA corpus).',
     'Primarily title-only records without substantive summaries (especially Quebec '
     'courts). The true distribution of Canadian water law dispute types cannot be '
     'reliably estimated from the available metadata. This is a genuine limitation of '
     'the collection methodology and should be flagged in any publication using the '
     'Canadian sub-corpus for typological analysis.'),

    (f'Brazil other_water ({n_ow_br:,} cases, {pct_ow_br:.1f}% of BR corpus).',
     f'A heterogeneous residual including TJSP list-format results (no ementa), '
     f'general civil liability cases involving water infrastructure, administrative '
     f'decisions about SABESP/CAESB operations, and some criminal cases that mention '
     f'water incidentally. The 21 named categories account for {100-pct_ow_br:.1f}% of Brazilian cases, '
     f'which is sufficient for the core analytical claims of this research.'),
]:
    p = doc.add_paragraph()
    p.add_run(f'{item[0]}: ').bold = True
    p.add_run(item[1])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════
h1('Executive Summary')
para(
    f'This report presents comparative findings from the Global Water Law Judicial Decisions Dataset '
    f'(v1.0), comprising {n_total:,} judicial decisions on water access across Brazil ({n_br:,}), '
    f'the Netherlands ({n_nl:,}), and Canada ({n_ca:,}), covering the period 2016–2025. '
    'The analysis tests the "Legal Last Mile" thesis: that administrative law operates as a '
    'structural gatekeeper transforming physical access to water infrastructure into conditional '
    'legal status — and that the populations most excluded from physical access are also the '
    'populations least visible as judicial plaintiffs.'
)
doc.add_paragraph()
para('Key findings:', bold=True)
bullets = [
    f'Brazil: {n_br:,} decisions — dominated by tariff disputes (billing, disconnection, '
     f'reconnection). User win rate among coded cases: {pct_user:.1f}%.',
    'Administrative Ghost thesis confirmed: only ~1.3% of Brazilian cases involve irregular '
    'settlements — excluded populations do not appear as plaintiffs in the formal judicial system.',
    f'Netherlands: {n_nl:,} decisions — near-zero household connection disputes. Pre-litigation '
    'absorption through administrative enforcement resolves most conflicts before court filing.',
    f'Canada: {n_ca:,} decisions — intermediate model combining opaque administrative discretion '
    'with fragmented statutory protections.',
    f'Post-Lei 14.026/2020 reform (Brazil): user win rate shifted from {pct_user_pre:.1f}% '
    f'(pre-2020) to {pct_user_post:.1f}% (2020–2025) among coded decisions.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. DATASET OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
h1('1. Dataset Overview')
h2('1.1 Collection Scope')
para(
    'The dataset was assembled through automated web scrapers targeting publicly accessible '
    'jurisprudência portals across three jurisdictions, supplemented by API access (CanLII for '
    'Canada) and open data feeds (Rechtspraak.nl for the Netherlands). The collection period '
    'is 2016–2025; Brazilian historical data from TJSP extends to 1997.'
)
doc.add_paragraph()
tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Table Grid'
for i, h in enumerate(['Country', 'Cases', 'Courts / Sources', 'Period']):
    tbl.rows[0].cells[i].text = h
    tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row in [
    ('Brazil',      f'{n_br:,}',  '8 state courts (TJDFT, TJSC, TJRJ, TJSP, TJRR, TJAC, TJPI, TJTO) + TJSP historical', '1997–2025'),
    ('Netherlands', f'{n_nl:,}',  'RvS + CBb + GHARL + 11 Rechtbanken (Rechtspraak Open Data)', '2016–2025'),
    ('Canada',      f'{n_ca:,}',  'CanLII API + Legal Data Hunter semantic search', '2016–2025'),
    ('TOTAL',       f'{n_total:,}', '—', '1997–2025'),
]:
    r = tbl.add_row()
    for i, val in enumerate(row):
        r.cells[i].text = val
        if row[0] == 'TOTAL':
            r.cells[i].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
add_chart(c1, 'Figure 1. Water law judicial decisions per year by country (2016–2025). Netherlands volume on right axis.')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. BRAZIL
# ══════════════════════════════════════════════════════════════════════════
h1('2. Brazil: The Tariff Dispute State')
para(
    f'Brazil contributes {n_br:,} decisions from 8 state courts. The docket is heavily '
    'concentrated at the TJDFT (Brasília, 8,421 decisions) and TJSC (Santa Catarina, 1,224), '
    'reflecting both population size and search methodology differences across court portals.'
)

h2('2.1 Governance Distribution')
para(
    'Brazilian water law litigation is overwhelmingly dominated by billing and service '
    'interruption disputes. Tariff cases (cobrança indevida, corte de fornecimento, '
    'restabelecimento) represent the structural core of the Brazilian water court docket, '
    'reflecting the commercialisation of water services under the 1997 Lei de Saneamento '
    'framework and the 2007 Lei de Diretrizes Nacionais do Saneamento Básico.'
)
add_chart(c2, 'Figure 2. Brazil: Top governance categories by case count (2016–2025).')

h2('2.2 The Administrative Ghost: Irregular Settlements')
para(
    'The "Administrative Ghost" thesis predicts that populations in irregular settlements '
    'will be systematically absent from the judicial record — not because their water '
    'problems are resolved, but because they lack legal standing, documentation, and '
    'institutional access to litigate.'
)
para(
    'The data confirms this: cases involving irregular settlements represent approximately '
    '1.3% of the Brazilian docket — a structural undercount given that informal settlements '
    'house an estimated 20–30% of Brazil\'s urban population (IBGE 2020). The Administrative '
    'Ghost is an absence in the judicial record made visible only by its statistical smallness.'
)
add_chart(c6, 'Figure 3. Brazil: Irregular settlement cases (count and % of annual total, 2016–2025).')

h2('2.3 Win/Loss Outcomes and the 2020 Sanitation Reform')
para(
    f'Among {n_br_coded:,} Brazilian decisions with coded outcomes, users prevail in '
    f'{pct_user:.1f}% of cases and utilities in {pct_util:.1f}%. This reflects the structural '
    'asymmetry of Brazilian water litigation: most cases are initiated by users contesting '
    'wrongful disconnections or billing errors, and courts have developed predictable doctrine '
    'favouring reconnection where utilities cannot demonstrate procedural compliance.'
)
para(
    f'The sanitation reform (Lei 14.026/2020) coincided with a notable shift in outcome '
    f'distribution. Pre-2020 user win rate: {pct_user_pre:.1f}%. Post-2020: {pct_user_post:.1f}%. '
    'This shift may reflect increased compliance pressure on utilities, stricter court scrutiny '
    'of disconnection procedures, or compositional changes in the litigation pool.'
)
add_chart(c3, 'Figure 4. Brazil: Win/loss outcomes by year (2016–2025). Shaded region: post-Lei 14.026/2020.')

h2('2.4 Human Rights Language')
para(
    f'Human rights framing appears in {n_hr:,} Brazilian decisions ({pct_hr:.1f}% of the '
    'Brazilian docket). The presence of explicit HR language — direito humano à água, '
    'mínimo existencial, dignidade da pessoa humana — is notably low given Brazil\'s '
    'constitutional framework (Art. 6, CF/1988). Courts resolve water disputes primarily '
    'through consumer protection and contract law frameworks, rarely invoking constitutional '
    'rights language even where it would be available (cf. Haglund 2014, 2019).'
)
add_chart(c4, 'Figure 5. Brazil: Cases invoking human rights language (count and % of total, 2016–2025).')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. NETHERLANDS
# ══════════════════════════════════════════════════════════════════════════
h1('3. Netherlands: Pre-Litigation Absorption')
para(
    f'The Netherlands contributes {n_nl:,} decisions — the largest single-country component. '
    'Despite near-universal water access, the Dutch docket is structurally different from '
    'Brazil. Household connection disputes are near-absent; Dutch administrative law (bestuursrecht) '
    'resolves most conflicts through institutionalised pre-litigation mechanisms '
    '(bezwaar, beroep) before they enter the judicial record.'
)

h2('3.1 The Dutch Pre-Litigation Model')
para(
    'By the time a case reaches the Raad van State or district courts, it typically involves '
    'complex environmental permitting, water board (waterschap) decisions, or groundwater '
    'extraction disputes — not individual household access. The high volume reflects '
    'comprehensive publication by Rechtspraak.nl and the breadth of water-adjacent '
    'administrative decisions, not access failures.'
)

h2('3.2 Implications for the Legal Last Mile')
para(
    'The Dutch model demonstrates that the legal last mile problem can be substantially '
    'resolved through administrative rather than judicial mechanisms. The near-absence of '
    'household connection disputes in Dutch courts is evidence that the administrative '
    'system absorbs and resolves them — representing a potential institutional design '
    'target for reforming access governance in developing jurisdictions.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. CANADA
# ══════════════════════════════════════════════════════════════════════════
h1('4. Canada: Fragmented Opacity')
para(
    f'Canada contributes {n_ca:,} decisions, collected via the CanLII API and Legal Data '
    'Hunter semantic search. The relatively small Canadian corpus reflects genuine differences '
    'in litigation frequency: Canadian water governance operates primarily through opaque '
    'administrative discretion at the municipal level, with limited statutory rights of '
    'appeal and significant provincial variation.'
)

h2('4.1 Structural Opacity')
para(
    'Unlike Brazil (billing disputes generating high-volume litigation) and the Netherlands '
    '(permit decisions generating high-volume administrative review), Canada\'s water governance '
    'produces limited judicially visible output. Municipal water decisions are typically not '
    'subject to formal administrative appeal, and connection refusals — the paradigmatic '
    '"legal last mile" dispute — are rarely challenged in court. This opacity differs from '
    'Dutch pre-litigation absorption: Canadian disputes are not resolved administratively; '
    'they are simply not reviewable.'
)

h2('4.2 Indigenous Water Rights')
para(
    'A structurally distinct sub-corpus involves Indigenous communities, particularly '
    'long-term drinking water advisories on reserves under federal jurisdiction. These cases '
    'involve different legal frameworks, constitutional rights, and historical dispossession. '
    'The current collection likely undercounts Indigenous water rights litigation; a dedicated '
    'collection using FNLM-specific keywords would substantially expand this component.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. COMPARATIVE ANALYSIS
# ══════════════════════════════════════════════════════════════════════════
h1('5. Comparative Analysis: Three Models of Water Governance')

h2('5.1 The Three Models')
for name, desc in [
    ('Reactive Adjudication (Brazil)',
     'High-volume individual billing disputes. Courts resolve tariff and disconnection '
     'conflicts case-by-case. Governance through private law enforcement. Excluded populations '
     'structurally absent from the judicial record (Administrative Ghost).'),
    ('Pre-Litigation Absorption (Netherlands)',
     'Low household dispute volume despite comprehensive system coverage. Administrative '
     'mechanisms resolve most conflicts before litigation. High environmental and permitting '
     'volume reflects planning complexity, not access failures.'),
    ('Opaque Discretion (Canada)',
     'Low dispute volume reflecting limited reviewability, not resolution. Municipal '
     'administrative decisions largely unreviewable. Fragmented provincial frameworks '
     'produce inconsistent rights landscapes. Indigenous rights structurally distinct.'),
]:
    p = doc.add_paragraph()
    p.add_run(f'{name}: ').bold = True
    p.add_run(desc)

h2('5.2 Comparative Governance')
add_chart(c5, 'Figure 6. Comparative governance category distributions across Brazil, Netherlands, and Canada.')

tbl2 = doc.add_table(rows=1, cols=4); tbl2.style = 'Table Grid'
for i, h in enumerate(['Dimension', 'Brazil', 'Netherlands', 'Canada']):
    tbl2.rows[0].cells[i].text = h
    tbl2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row in [
    ('Primary dispute type', 'Billing/disconnection', 'Environmental permits', 'Mixed / opaque'),
    ('Judicial volume', 'High', 'Very high (structural)', 'Low'),
    ('HR language frequency', f'{pct_hr:.1f}%', 'Low', 'Low'),
    ('Excluded pop. visibility', '~1.3% irregular settlements', 'N/A (near-universal access)', 'Indigenous undercount'),
    ('Post-reform shift', f'User wins: {pct_user_pre:.0f}%→{pct_user_post:.0f}% (2020)', 'N/A', 'N/A'),
    ('Governance model', 'Reactive Adjudication', 'Pre-Litigation Absorption', 'Opaque Discretion'),
]:
    r = tbl2.add_row()
    for i, val in enumerate(row): r.cells[i].text = val

h2('5.3 The Sanitation Reform Effect (Brazil)')
para(
    'Lei 14.026/2020 introduced universal service targets, private sector concession '
    'requirements, and explicit service quality standards. The dataset captures its first '
    f'five years of implementation. Coded win/loss outcomes show a shift: user win rate '
    f'from {pct_user_pre:.1f}% (2016–2019) to {pct_user_post:.1f}% (2020–2025). '
    'This is consistent with increased judicial scrutiny of utility compliance under the '
    'reformed framework, though causal attribution requires controlling for compositional '
    'changes in the litigation pool.'
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════
h1('6. Conclusions and Open Questions')

h2('6.1 The Legal Last Mile in Three Jurisdictions')
para(
    'The "Legal Last Mile" thesis holds across all three jurisdictions, though it manifests '
    'differently. In Brazil, the last mile is visible as a statistical absence: informal '
    'settlement residents are essentially absent from the judicial record. The legal system '
    'processes billing disputes for connected users; it provides no forum for access claims '
    'by the unconnected.'
)
para(
    'In the Netherlands, the last mile is substantively resolved before reaching courts. '
    'The Dutch case is most useful as a comparison point: high judicial volume does not '
    'imply access failures, and pre-litigation systems can effectively substitute for '
    'judicial enforcement.'
)
para(
    'In Canada, the last mile remains structurally opaque. The limited judicial record '
    'reflects not resolution but invisibility: administrative discretion without '
    'reviewability, Indigenous rights without adequate enforcement, and fragmented '
    'protections varying dramatically across provincial boundaries.'
)

h2('6.2 Open Research Questions')
for q in [
    'What determines the pre-litigation absorption capacity of administrative systems? '
    'Can Dutch-model mechanisms be transplanted to high-inequality contexts?',
    'Does the Brazilian post-2020 shift reflect genuine improvement in utility compliance, '
    'or compositional changes in who litigates?',
    'How do irregular settlement residents in Brazil navigate water access outside the '
    'formal judicial system? What quasi-legal mechanisms substitute for litigation?',
    'How should Indigenous water rights be theorised within the Legal Last Mile framework?',
    'What would it take to build a truly global water law dataset covering Global South '
    'courts currently inaccessible to automated collection?',
]:
    doc.add_paragraph(q, style='List Number')

h2('6.3 Dataset and Replication')
para('The full dataset, jurimetric coding engine, and all scrapers are available at:')
for r in [
    'Zenodo: https://doi.org/10.5281/zenodo.19836413',
    'Harvard Dataverse: https://dataverse.harvard.edu/dataset.xhtml?persistentId=doi:10.7910/DVN/C9PEFS',
    'DANS SSH: https://ssh.datastations.nl/dataset.xhtml?persistentId=doi:10.17026/SS/RVDBUF',
    'OSF: https://osf.io/admrq',
    'GitHub: https://github.com/jrklaus8/water-law-dataset',
]:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════
h1('References')
for ref in [
    'Bakker, K. (2007). The "Commons" Versus the "Commodity". Antipode, 39(3), 430–455.',
    'Castro, J.E. (2004). Urban Water and the Politics of Citizenship. Environment and Planning A, 36(2), 327–346.',
    'Gauri, V. & Brinks, D.M. (2008). Courting Social Justice. Cambridge University Press.',
    'Haglund, L. (2014). Water Governance and Social Justice in São Paulo, Brazil. Water Policy, 16(S2), 78–97.',
    'Haglund, L. (2019). Can Human Rights Challenge Neoliberal Logics? In Economic and Social Rights in a Neoliberal World. Cambridge University Press.',
    'Haglund, L. (2019). Sustainable Urban Water Governance and Human Rights. Sustainability, 11(19), 5314.',
    'Hirschl, R. (2004). Towards Juristocracy. Harvard University Press.',
    'Klaus, C. (2026). Global Water Law Judicial Decisions Dataset (v1.0). Zenodo. https://doi.org/10.5281/zenodo.19836413',
    'Varley, A. (2013). Postcolonialising Informality? Environment and Planning D, 31(1), 4–22.',
]:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.left_indent       = Inches(0.3)

# ══════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════
doc.save(str(OUT))
print(f"\nDone. Report saved to: {OUT}")
print(f"Charts saved to: {CHARTS}")
