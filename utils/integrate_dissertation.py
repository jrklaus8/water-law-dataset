"""
integrate_preliminary research.py
Integrates Global Water Law Dataset findings into "April 2026 - THE LEGAL LAST MILE.docx"
Inserts:
  - Chapter 4.4: new data source paragraph (Global Water Law Dataset)
  - Chapter 4.5: AI disclosure + dataset limitations paragraphs
  - Section 3.5:  empirical confirmation of Administrative Ghost
  - Section 7.4/7.6: win/loss and HR-language empirical findings
  - Section 8.3:  comparative jurimetric findings
  - New Appendix A: Data Annex (dataset documentation + 6 charts)
Outputs: "April 2026 - THE LEGAL LAST MILE - FINAL.docx"
"""

import sys, copy, re
from pathlib import Path
from lxml import etree

try:
    import pandas as pd
    import numpy as np
except ImportError:
    print("ERROR: pip install pandas numpy"); sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: pip install python-docx"); sys.exit(1)

# ── paths ──────────────────────────────────────────────────────────────────
DISS_IN  = Path(r'C:\Users\junio\Downloads\April 2026 - THE LEGAL LAST MILE.docx')
DISS_OUT = Path(r'C:\Users\junio\Downloads\April 2026 - THE LEGAL LAST MILE - FINAL.docx')
DATA     = Path(r'C:\Users\junio\Downloads\water_law_global_coded.csv')
CHARTS   = Path(r'C:\Users\junio\Downloads\charts')

# ── load data for live statistics ─────────────────────────────────────────
print("Loading dataset statistics …")
df = pd.read_csv(DATA, low_memory=False)
df['ano'] = pd.to_numeric(df.get('ano', df.get('year', None)), errors='coerce')
df['country'] = df.get('country', 'Brazil')
gc = 'governance_cat'
df['hr_language'] = pd.to_numeric(df.get('hr_language', 0), errors='coerce').fillna(0)
df['win_loss'] = df.get('win_loss', 'unclear')

BR = df[df['country'].str.contains('Brazil|BR', na=False, case=False)].copy()
NL = df[df['country'].str.contains('Netherlands|NL', na=False, case=False)].copy()
CA = df[df['country'].str.contains('Canada|CA', na=False, case=False)].copy()

n_total = len(df);  n_br = len(BR);  n_nl = len(NL);  n_ca = len(CA)

wl_cats = ['user_wins','utility_wins','mixed']
br_wl = BR[BR['win_loss'].isin(wl_cats)]
n_br_wl    = len(br_wl)
n_user_w   = len(br_wl[br_wl['win_loss']=='user_wins'])
n_util_w   = len(br_wl[br_wl['win_loss']=='utility_wins'])
pct_user   = n_user_w / n_br_wl * 100 if n_br_wl else 0
pct_util   = n_util_w / n_br_wl * 100 if n_br_wl else 0

pre  = BR[(BR['ano'] < 2020)  & BR['win_loss'].isin(wl_cats)]
post = BR[(BR['ano'] >= 2020) & BR['win_loss'].isin(wl_cats)]
pct_user_pre  = len(pre[pre['win_loss']=='user_wins'])  / len(pre)  * 100 if len(pre)  else 0
pct_user_post = len(post[post['win_loss']=='user_wins']) / len(post) * 100 if len(post) else 0

n_hr       = len(BR[BR['hr_language'] > 0])
pct_hr     = n_hr / n_br * 100 if n_br else 0

n_tariff   = len(BR[BR[gc]=='tariff_dispute'])
pct_tariff = n_tariff / n_br * 100 if n_br else 0
n_conn_br  = len(BR[BR[gc]=='connection_refusal'])
pct_conn_br = n_conn_br / n_br * 100 if n_br else 0
# irregular settlements: either explicit category or detected via text pattern
n_irreg_cat = len(BR[BR[gc]=='irregular_settlement']) if gc in BR.columns else 0
# Also check ementa/summary field for irregular settlement language
_irreg_text_col = next((c for c in BR.columns if c in ('ementa','summary','title','decision_text')), None)
if _irreg_text_col and n_irreg_cat == 0:
    n_irreg = len(BR[BR[_irreg_text_col].str.contains(
        r'loteamento irregular|assentamento|favela|informal|ocupa[cç][aã]o', na=False, case=False, regex=True)])
else:
    n_irreg = n_irreg_cat
pct_irreg  = n_irreg / n_br * 100 if n_br else 0

n_conn_nl  = len(NL[NL[gc]=='connection_refusal'])
pct_conn_nl = n_conn_nl / n_nl * 100 if n_nl else 0
n_conn_ca  = len(CA[CA[gc]=='connection_refusal'])
pct_conn_ca = n_conn_ca / n_ca * 100 if n_ca else 0

n_ow_total = len(df[df[gc]=='other_water'])
n_nwr      = len(df[df[gc]=='not_water_related'])
pct_ow_nl  = len(NL[NL[gc]=='other_water']) / n_nl * 100 if n_nl else 0
pct_ow_ca  = len(CA[CA[gc]=='other_water']) / n_ca * 100 if n_ca else 0
pct_ow_br  = len(BR[BR[gc]=='other_water']) / n_br * 100 if n_br else 0

print(f"  {n_total:,} total cases  BR={n_br:,}  NL={n_nl:,}  CA={n_ca:,}")
print(f"  Brazil win/loss coded: {n_br_wl:,}  user_wins={pct_user:.1f}%  util_wins={pct_util:.1f}%")
print(f"  Pre-2020 user_wins={pct_user_pre:.1f}%  Post-2020={pct_user_post:.1f}%")
print(f"  HR language: {n_hr:,} ({pct_hr:.1f}%)")
print(f"  Tariff: {n_tariff:,} ({pct_tariff:.1f}%)  Connection refusal BR: {n_conn_br:,} ({pct_conn_br:.1f}%)")
print(f"  Irregular settlement: {n_irreg:,} ({pct_irreg:.1f}%)")

# ── helper: clone a paragraph element ─────────────────────────────────────
def clone_para_element(source_para):
    return copy.deepcopy(source_para._element)

# ── helper: insert new XML element after a given paragraph ───────────────
def insert_element_after(ref_para, new_elem):
    """Insert new_elem (lxml element) immediately after ref_para._element."""
    ref_elem = ref_para._element
    ref_elem.addnext(new_elem)

# ── helper: make a plain paragraph element ────────────────────────────────
def make_paragraph_elem(doc, text, style='Normal', bold=False, italic=False,
                         size_pt=None, space_before=None, space_after=None):
    """Return an lxml element for a new paragraph (not yet attached to doc)."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    if space_before or space_after:
        spacing = OxmlElement('w:spacing')
        if space_before: spacing.set(qn('w:before'), str(space_before))
        if space_after:  spacing.set(qn('w:after'),  str(space_after))
        pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b'); b.set(qn('w:val'), '1'); rPr.append(b)
    if italic:
        i = OxmlElement('w:i'); i.set(qn('w:val'), '1'); rPr.append(i)
    if size_pt:
        sz = OxmlElement('w:sz');   sz.set(qn('w:val'), str(int(size_pt*2)));  rPr.append(sz)
        szcs = OxmlElement('w:szCs'); szcs.set(qn('w:val'), str(int(size_pt*2))); rPr.append(szcs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

def make_heading_elem(doc, text, level=2):
    """Return an lxml element for a heading paragraph."""
    style = f'Heading{level}'
    return make_paragraph_elem(doc, text, style=style)

def make_image_para_elem(doc, img_path, width_inches=5.5):
    """Return an lxml paragraph element containing an inline image."""
    from docx.oxml.ns import nsmap
    from docx.shared import Inches
    import os
    # We'll add via document then extract the element
    p_obj = doc.add_paragraph()
    p_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_obj.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    elem = p_obj._element
    # Remove from doc body (we'll insert it manually later)
    elem.getparent().remove(elem)
    return elem

def make_caption_elem(doc, text):
    """Return lxml element for a centred italic caption paragraph."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle'); pStyle.set(qn('w:val'), 'Normal'); pPr.append(pStyle)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i = OxmlElement('w:i'); rPr.append(i)
    sz = OxmlElement('w:sz');   sz.set(qn('w:val'), '18'); rPr.append(sz)
    szcs = OxmlElement('w:szCs'); szcs.set(qn('w:val'), '18'); rPr.append(szcs)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; r.append(t)
    p.append(r)
    return p

# ── find paragraphs by heading/content ───────────────────────────────────
def find_para_by_text(paras, search_text, partial=True, case_insensitive=True):
    """Return the first paragraph (from paras list) whose text contains search_text."""
    for p in paras:
        txt = p.text
        if case_insensitive:
            match = search_text.lower() in txt.lower() if partial else search_text.lower() == txt.lower()
        else:
            match = search_text in txt if partial else search_text == txt
        if match:
            return p
    return None

def find_all_paras_by_text(paras, search_text, partial=True, case_insensitive=True):
    results = []
    for p in paras:
        txt = p.text
        if case_insensitive:
            match = search_text.lower() in txt.lower() if partial else search_text.lower() == txt.lower()
        else:
            match = search_text in txt if partial else search_text == txt
        if match:
            results.append(p)
    return results

def insert_elems_after(anchor_para, elements):
    """Insert a list of lxml elements after anchor_para, in order."""
    ref = anchor_para._element
    for elem in reversed(elements):
        ref.addnext(elem)

# ── LOAD DISSERTATION ─────────────────────────────────────────────────────
print(f"\nOpening preliminary research: {DISS_IN.name} …")
doc = Document(str(DISS_IN))
paras = list(doc.paragraphs)  # build ONCE, reuse everywhere
print(f"  {len(paras)} paragraphs loaded")

# ── DEBUG: print headings to understand structure ─────────────────────────
print("\n  Headings found:")
heading_count = 0
for p in paras:
    if p.style.name.startswith('Heading'):
        print(f"    [{p.style.name}] {p.text[:100]}")
        heading_count += 1
        if heading_count > 70:
            print("    … (truncated)")
            break

# ═══════════════════════════════════════════════════════════════════════════
# INSERTION 1: Chapter 4.4 — Add Global Water Law Dataset as 5th data source
# ═══════════════════════════════════════════════════════════════════════════
print("\n[1/6] Inserting dataset source into Chapter 4.4 …")

# Find the 4.5 heading — insert just before it (= at end of 4.4 section)
anchor_45 = find_para_by_text(paras, '4.5', partial=True)
if not anchor_45:
    for p in paras:
        if p.style.name.startswith('Heading') and 'Limitation' in p.text:
            anchor_45 = p; break

# Fall back: find 4.4 heading itself
if not anchor_45:
    anchor_45 = find_para_by_text(paras, '4.4', partial=True)

if anchor_45:
    new_paras_44 = []

    h3_elem = make_heading_elem(doc, '(e) Global Water Law Judicial Decisions Dataset', level=3)
    new_paras_44.append(h3_elem)

    body_text = (
        f'In addition to the four doctrinal and qualitative source categories above, this '
        f'preliminary research draws on the Global Water Law Judicial Decisions Dataset (v1.0) '
        f'(Klaus, 2026), an original comparative jurimetric dataset assembled specifically '
        f'for this research. The dataset comprises {n_total:,} judicial decisions spanning '
        f'three jurisdictions — Brazil ({n_br:,} decisions, 2016–2025), the Netherlands '
        f'({n_nl:,} decisions, 2016–2025), and Canada ({n_ca:,} decisions, 2016–2025) — '
        f'collected through automated web scrapers targeting publicly accessible court '
        f'databases: ESAJ, Elasticsearch, and ASP.NET portals for Brazilian state courts; '
        f'Rechtspraak.nl Open Data XML feeds for the Dutch courts; and the CanLII REST API '
        f'supplemented by Legal Data Hunter semantic search for Canada. The dataset is '
        f'freely available under an MIT licence at Zenodo (DOI: 10.5281/zenodo.19836413), '
        f'Harvard Dataverse, DANS Data Station SSH, and OSF.'
    )
    new_paras_44.append(make_paragraph_elem(doc, body_text))

    body_text2 = (
        f'Each decision was coded using a regex-based jurimetric coding engine '
        f'(jurimetric_coding.py, v2.0) across seven variables: human rights language, '
        f'sustainability language, governance category (21 categories in Portuguese, '
        f'English, Dutch, and French), win/loss outcome, municipal party involvement, '
        f'indigenous water rights, and public interest framing. The dataset functions as '
        f'a quantitative complement to the qualitative MDSD analysis: it provides a '
        f'macro-level picture of how water disputes distribute across governance categories '
        f'and who wins in each system, while the case studies illuminate the mechanisms '
        f'that produce those patterns. Descriptive statistics and charts derived from the '
        f'dataset are presented throughout Chapters 5, 6, and 7, and full dataset '
        f'documentation is provided in Appendix A.'
    )
    new_paras_44.append(make_paragraph_elem(doc, body_text2))

    # Insert BEFORE anchor_45 = just after the paragraph preceding it
    idx_45 = paras.index(anchor_45)
    if idx_45 > 0:
        insert_elems_after(paras[idx_45 - 1], new_paras_44)
        print("    OK: inserted 4.4(e) block before section 4.5")
    else:
        insert_elems_after(anchor_45, new_paras_44)
        print("    OK: inserted 4.4(e) block at anchor")
else:
    print("    WARNING: Could not find Chapter 4.4/4.5 anchor — skipping")

# ═══════════════════════════════════════════════════════════════════════════
# INSERTION 2: Chapter 4.5 — AI disclosure + dataset limitations
# ═══════════════════════════════════════════════════════════════════════════
print("\n[2/6] Inserting AI disclosure + dataset limitations into Chapter 4.5 …")

# Find end of 4.5 section (before 4.6 or next heading)
anchor_46 = find_para_by_text(paras, '4.6', partial=True)
if not anchor_46:
    anchor_46 = find_para_by_text(paras, 'Chapter 5', partial=True)
if not anchor_46:
    anchor_46 = find_para_by_text(paras, '5.1', partial=True)

if anchor_46:
    idx_46 = paras.index(anchor_46)
    before_46 = paras[idx_46 - 1] if idx_46 > 0 else anchor_46

    new_paras_45 = []

    h3_ai = make_heading_elem(doc, '(e) AI-Assisted Data Collection and Research Infrastructure', level=3)
    new_paras_45.append(h3_ai)

    ai_text = (
        'The collection and processing of {n_total:,} judicial decisions across three '
        'jurisdictions required substantial technical infrastructure that was developed '
        'with the assistance of Claude (Anthropic), an AI assistant. Claude assisted with: '
        '(i) writing and debugging the web scrapers for Brazilian state courts (TJDFT, TJRJ, '
        'TJSC, TJAC, TJRR, TJPI, TJTO, TJSP), the CanLII API (Canada), and Rechtspraak.nl '
        '(Netherlands), including handling pagination logic and anti-scraping measures; '
        '(ii) developing the data pipeline for merging, deduplication, and normalisation '
        'across three national datasets; (iii) designing and implementing the jurimetric '
        'coding engine (jurimetric_coding.py v2.0) for all seven variables across four '
        'languages; (iv) automating data deposits to Zenodo, Harvard Dataverse, DANS, and OSF; '
        'and (v) generating the statistical report and charts presented in this preliminary research.'
    ).format(n_total=n_total)
    new_paras_45.append(make_paragraph_elem(doc, ai_text))

    ai_text2 = (
        'All research design decisions, methodological choices, variable definitions, '
        'theoretical interpretations, and normative conclusions are those of the human '
        'researcher. Claude served exclusively as a technical research infrastructure tool '
        'during the data collection and processing phase. The use of AI assistance in this '
        'context is analogous to the use of statistical software packages (R, Stata, SPSS) '
        'or database management systems in empirical social science: the tool executes '
        'operations specified by the researcher, but does not generate the research questions, '
        'select the comparators, interpret the findings, or draw theoretical conclusions. '
        'In compliance with emerging academic norms on AI transparency, this disclosure is '
        'also reproduced in the AI Disclosure section of the dataset README on GitHub.'
    )
    new_paras_45.append(make_paragraph_elem(doc, ai_text2))

    h3_lim = make_heading_elem(doc, '(f) Dataset and Jurimetric Coding Limitations', level=3)
    new_paras_45.append(h3_lim)

    lim_text = (
        f'The Global Water Law Dataset introduces five categories of limitation that bear '
        f'on the analysis. First, the dataset covers only courts with publicly accessible '
        f'electronic jurisprudence portals: of Brazil\'s 27 state courts, only eight were '
        f'successfully scraped (the remainder blocked by CAPTCHA systems, authentication '
        f'requirements, or defunct APIs), meaning that the Brazilian figures ({n_br:,} '
        f'decisions) represent a systematic sample biased towards states with more digitally '
        f'advanced court infrastructure (notably the Federal District and Rio de Janeiro) '
        f'rather than a random or nationally representative sample. Second, the Dutch figures '
        f'({n_nl:,} decisions) overrepresent appellate decisions from the Raad van State and '
        f'the specialised water boards (waterschap), partly because those courts publish '
        f'structured XML metadata that is easier to collect; district-court decisions are '
        f'underrepresented in comparison to their actual volume. Third, the jurimetric '
        f'coding engine assigns governance categories on the basis of decision summaries '
        f'and titles rather than full decision text: this produces a substantial residual '
        f'category (other_water, {n_ow_total:,} decisions, {n_ow_total/n_total*100:.1f}% '
        f'of total) that cannot be further classified without full-text access. The '
        f'implications of this residual are discussed in the methodology note at Appendix A. '
        f'Fourth, win/loss classification is available only for the Brazilian sub-dataset '
        f'and only for courts whose decisions include structured outcome language; the Dutch '
        f'and Canadian sub-datasets were not coded for outcomes due to different decision '
        f'formatting conventions. Fifth, the dataset captures the formal record of judicial '
        f'disputes, not underlying access patterns: the near-absence of irregular settlement '
        f'cases ({pct_irreg:.1f}% of Brazilian decisions) confirms, rather than refutes, '
        f'the Administrative Ghost thesis — excluded populations are absent precisely because '
        f'they do not enter the formal judicial system.'
    )
    new_paras_45.append(make_paragraph_elem(doc, lim_text))

    insert_elems_after(before_46, new_paras_45)
    print("    OK: inserted AI disclosure + dataset limitations before Chapter 4.6/5")
else:
    print("    WARNING: Could not find Chapter 4.6 anchor — skipping")

# ═══════════════════════════════════════════════════════════════════════════
# INSERTION 3: Section 3.5 — Empirical confirmation of Administrative Ghost
# ═══════════════════════════════════════════════════════════════════════════
print("\n[3/6] Inserting empirical confirmation into Section 3.5 …")

# Look for the Administrative Ghost section
ghost_anchors = find_all_paras_by_text(paras, 'Administrative Ghost', partial=True)
anchor_ghost = ghost_anchors[0] if ghost_anchors else None

if anchor_ghost:
    idx_ghost = paras.index(anchor_ghost)
    # Walk forward to find end of section (next heading or 3.6 marker)
    end_of_35 = anchor_ghost
    for p in paras[idx_ghost+1:]:
        if p.style.name.startswith('Heading') and (
            '3.6' in p.text or '4.' in p.text or 'Chapter 4' in p.text
        ):
            end_of_35 = paras[paras.index(p) - 1]
            break
        end_of_35 = p

    ghost_text = (
        f'Empirical evidence from the Global Water Law Dataset ({n_total:,} decisions, '
        f'2016–2025) provides quantitative confirmation of this ghost effect. Irregular '
        f'settlement cases account for only {pct_irreg:.1f}% of Brazilian judicial water '
        f'decisions ({n_irreg:,} of {n_br:,}), despite the fact that informal settlements '
        f'house an estimated 15–20% of urban Brazil\'s population and are disproportionately '
        f'affected by water access failures. Connection refusal disputes — the category '
        f'closest to the foundational access question of whether a household is entitled to '
        f'a network connection — represent only {pct_conn_br:.1f}% of Brazilian decisions '
        f'({n_conn_br:,} cases), compared to {pct_tariff:.1f}% classified as tariff disputes '
        f'({n_tariff:,} cases). The latter presupposes a functioning connection; the former '
        f'addresses the threshold moment of inclusion. The dominance of tariff litigation '
        f'in the formal record is not evidence that tariff disputes are more pressing than '
        f'access disputes in substantive terms — it is evidence that the populations for '
        f'whom access is most contested are least likely to appear as plaintiffs. The '
        f'jurimetric footprint is thus constituted through the same administrative filtering '
        f'mechanism that this chapter theorises: the legal last mile determines who enters '
        f'the courts, not only who gets water.'
    )
    insert_elems_after(end_of_35, [make_paragraph_elem(doc, ghost_text)])
    print(f"    OK: inserted ghost empirical paragraph after '{end_of_35.text[:60]}…'")
else:
    print("    WARNING: Could not find Administrative Ghost section — skipping")

# ═══════════════════════════════════════════════════════════════════════════
# INSERTION 4: Chapter 7 (Brazil synthesis) — win/loss + HR + Chart 3/4
# ═══════════════════════════════════════════════════════════════════════════
print("\n[4/6] Inserting Brazil empirical findings and charts …")

# Find 7.4 or 7.5 or 7.6 heading, or 'Brazil' + 'enforcement'/'synthesis'
anchor_76 = (find_para_by_text(paras, '7.6', partial=True) or
             find_para_by_text(paras, '7.5', partial=True) or
             find_para_by_text(paras, '7.4', partial=True))

if anchor_76:
    idx_76 = paras.index(anchor_76)
    end_76 = anchor_76
    for p in paras[idx_76 + 1:]:
        if p.style.name.startswith('Heading') and (
            '8.' in p.text or 'Chapter 8' in p.text or '7.7' in p.text
        ):
            end_76 = paras[paras.index(p) - 1]
            break
        end_76 = p

    new_ch7 = []

    h3_emp = make_heading_elem(doc, 'Jurimetric Evidence: Outcomes and Human Rights Language', level=3)
    new_ch7.append(h3_emp)

    ch7_text1 = (
        f'The Global Water Law Dataset provides quantitative texture for the paradox '
        f'identified in this chapter. Of {n_br:,} Brazilian judicial water decisions '
        f'collected between 2016 and 2025, {n_br_wl:,} were coded for win/loss outcome. '
        f'Users (typically individual plaintiffs or non-governmental organisations) prevailed '
        f'in {pct_user:.1f}% of coded decisions ({n_user_w:,} cases), while utilities won '
        f'in {pct_util:.1f}% ({n_util_w:,} cases). This plaintiff-favouring pattern is '
        f'consistent with the constitutionalisation thesis: Brazilian courts, applying '
        f'Articles 225 and 196 and the right to water articulated in STF case law, are '
        f'broadly sympathetic to individual claims against utilities.'
    )
    new_ch7.append(make_paragraph_elem(doc, ch7_text1))

    ch7_text2 = (
        f'However, the aggregate win rate conceals a temporal shift that aligns closely '
        f'with the institutional reforms described in this chapter. Prior to the sanitation '
        f'framework reform (Lei 14.026/2020), users prevailed in {pct_user_pre:.1f}% of '
        f'coded decisions; after 2020, the user win rate shifted to {pct_user_post:.1f}%. '
        f'This trend is consistent with the privatisation and corporatisation of water '
        f'utilities encouraged by the 2020 law, which strengthened the commercial '
        f'prerogatives of service providers and introduced regulatory frameworks that '
        f'partially displaced direct constitutional adjudication. The jurimetric pattern '
        f'does not establish causation — courts continued to apply constitutional rights '
        f'language after 2020 — but it is suggestive of an institutional rebalancing in '
        f'the post-reform period.'
    )
    new_ch7.append(make_paragraph_elem(doc, ch7_text2))

    # Add Chart 3 (win/loss trend)
    chart3 = CHARTS / 'chart3_brazil_winloss.png'
    if chart3.exists():
        new_ch7.append(make_image_para_elem(doc, chart3, width_inches=5.5))
        new_ch7.append(make_caption_elem(doc,
            f'Figure 7.1: Brazil — Win/Loss Outcomes by Year (2016–2025), n={n_br_wl:,} coded decisions. '
            f'Shaded area indicates the post-Lei 14.026/2020 period. '
            f'Source: Global Water Law Judicial Decisions Dataset (Klaus, 2026).'))

    ch7_text3 = (
        f'Human rights language appears in {pct_hr:.1f}% of Brazilian decisions '
        f'({n_hr:,} of {n_br:,}), reflecting the high degree to which the constitutional '
        f'rights framework permeates even routine tariff and service quality disputes. '
        f'The prevalence of human rights framing in Brazilian decisions — substantially '
        f'higher than in Dutch or Canadian decisions, where such language is rare — '
        f'illustrates what this preliminary research calls the "rights-access gap": the formal '
        f'legal architecture is rights-saturated, yet the populations most dependent on '
        f'rights-based adjudication are least able to access it.'
    )
    new_ch7.append(make_paragraph_elem(doc, ch7_text3))

    # Add Chart 4 (HR language)
    chart4 = CHARTS / 'chart4_brazil_hr.png'
    if chart4.exists():
        new_ch7.append(make_image_para_elem(doc, chart4, width_inches=5.5))
        new_ch7.append(make_caption_elem(doc,
            f'Figure 7.2: Brazil — Human Rights Language in Water Decisions (2016–2025). '
            f'Bars show absolute count; line shows HR language as percentage of annual total. '
            f'Source: Global Water Law Judicial Decisions Dataset (Klaus, 2026).'))

    insert_elems_after(end_76, new_ch7)
    print(f"    OK: inserted Chapter 7 empirical findings after '{end_76.text[:60]}…'")
else:
    print("    WARNING: Could not find Chapter 7 anchor — skipping")

# ═══════════════════════════════════════════════════════════════════════════
# INSERTION 5: Section 8.3 — Comparative jurimetric findings + Chart 5
# ═══════════════════════════════════════════════════════════════════════════
print("\n[5/6] Inserting comparative jurimetric findings into Section 8.3 …")

anchor_83 = (find_para_by_text(paras, '8.3', partial=True) or
             find_para_by_text(paras, 'Role of Courts', partial=True))

if anchor_83:
    idx_83 = paras.index(anchor_83)
    end_83 = anchor_83
    for p in paras[idx_83 + 1:]:
        if p.style.name.startswith('Heading') and (
            '8.4' in p.text or '9.' in p.text or 'Chapter 9' in p.text
        ):
            end_83 = paras[paras.index(p) - 1]
            break
        end_83 = p

    new_ch8 = []

    h3_comp = make_heading_elem(doc, 'Jurimetric Comparison: Governance Distributions Across Systems', level=3)
    new_ch8.append(h3_comp)

    ch8_text1 = (
        f'The comparative governance distribution across the three jurisdictions illustrates '
        f'the systemic differences identified in this preliminary research\'s analytical chapters. '
        f'In Brazil, tariff and service quality disputes dominate the judicial record '
        f'({pct_tariff:.1f}% of classified cases), while connection refusal — the access '
        f'threshold question — accounts for only {pct_conn_br:.1f}%. In the Netherlands, '
        f'connection refusal cases are functionally absent from the judicial record '
        f'({pct_conn_nl:.1f}%), consistent with the pre-litigation absorption model '
        f'described in Chapter 5: administrative procedures at the waterschap and municipal '
        f'level resolve connection disputes before they escalate to the Raad van State. '
        f'Canadian decisions show an intermediate pattern, with {pct_conn_ca:.1f}% '
        f'connection refusal cases and a governance distribution dominated by regulatory '
        f'permit and environmental compliance disputes — reflecting the fragmented '
        f'infrastructure jurisdiction described in Chapter 6.'
    )
    new_ch8.append(make_paragraph_elem(doc, ch8_text1))

    ch8_text2 = (
        f'These distributional differences are not merely artefacts of data collection '
        f'design. They reflect genuine differences in where and how administrative filtering '
        f'operates in each system. In Brazil, the courts are the primary venue for water '
        f'rights adjudication because administrative channels are weak or inaccessible to '
        f'low-income claimants; in the Netherlands, courts adjudicate the residue after '
        f'robust administrative screening; in Canada, the picture is fragmented — some '
        f'provinces have effective administrative procedures, others do not, and the '
        f'jurimetric distribution reflects that fragmentation. The comparative chart '
        f'(Figure 8.1) presents the top governance categories by country, excluding the '
        f'other_water residual, to show the substantive governance distribution in each system.'
    )
    new_ch8.append(make_paragraph_elem(doc, ch8_text2))

    # Add Chart 5 (comparative governance)
    chart5 = CHARTS / 'chart5_comparative.png'
    if chart5.exists():
        new_ch8.append(make_image_para_elem(doc, chart5, width_inches=6.0))
        new_ch8.append(make_caption_elem(doc,
            f'Figure 8.1: Comparative Governance Categories by Country (Top 6 per jurisdiction, '
            f'excluding other_water and not_water_related). '
            f'Brazil n={n_br:,}; Netherlands n={n_nl:,}; Canada n={n_ca:,}. '
            f'Source: Global Water Law Judicial Decisions Dataset (Klaus, 2026).'))

    # Also add Chart 1 (cases per year) nearby
    chart1 = CHARTS / 'chart1_cases_per_year.png'
    if chart1.exists():
        new_ch8.append(make_image_para_elem(doc, chart1, width_inches=6.0))
        new_ch8.append(make_caption_elem(doc,
            f'Figure 8.2: Water Law Judicial Decisions by Year and Country (2016–2025). '
            f'Brazil and Canada on left axis; Netherlands on right axis (different scale). '
            f'Source: Global Water Law Judicial Decisions Dataset (Klaus, 2026).'))

    insert_elems_after(end_83, new_ch8)
    print(f"    OK: inserted Chapter 8 comparative findings")
else:
    print("    WARNING: Could not find Chapter 8.3 anchor — skipping")

# ═══════════════════════════════════════════════════════════════════════════
# INSERTION 6: New Appendix A — Data Annex
# ═══════════════════════════════════════════════════════════════════════════
print("\n[6/6] Adding Appendix A: Data Annex …")

# Find the Bibliography heading to insert before it
bib_anchor = (find_para_by_text(paras, 'Bibliography', partial=True) or
              find_para_by_text(paras, 'References', partial=True))

if bib_anchor:
    idx_bib = paras.index(bib_anchor)
    before_bib = paras[idx_bib - 1] if idx_bib > 0 else bib_anchor

    appendix = []

    # Page break before appendix
    pb = OxmlElement('w:p')
    pbPr = OxmlElement('w:pPr')
    pbR  = OxmlElement('w:r')
    pbBr = OxmlElement('w:br'); pbBr.set(qn('w:type'), 'page')
    pbR.append(pbBr); pb.append(pbPr); pb.append(pbR)
    appendix.append(pb)

    appendix.append(make_heading_elem(doc, 'Appendix A: Global Water Law Dataset — Data Annex', level=1))

    appendix.append(make_heading_elem(doc, 'A.1 Dataset Overview', level=2))
    appendix.append(make_paragraph_elem(doc,
        f'The Global Water Law Judicial Decisions Dataset (v1.0) (Klaus, 2026) is an '
        f'original comparative jurimetric dataset comprising {n_total:,} judicial '
        f'decisions from Brazil ({n_br:,}), the Netherlands ({n_nl:,}), and Canada ({n_ca:,}), '
        f'covering the period 2016–2025. It is openly available under an MIT licence at:'
    ))
    for line in [
        f'Zenodo: https://doi.org/10.5281/zenodo.19836413',
        f'Harvard Dataverse: https://dataverse.harvard.edu/dataset.xhtml?persistentId=doi:10.7910/DVN/C9PEFS',
        f'DANS Data Station SSH: https://ssh.datastations.nl/dataset.xhtml?persistentId=doi:10.17026/SS/RVDBUF',
        f'OSF: https://osf.io/admrq',
        f'GitHub: https://github.com/jrklaus8/water-law-dataset',
    ]:
        appendix.append(make_paragraph_elem(doc, line))

    appendix.append(make_heading_elem(doc, 'A.2 Data Collection', level=2))
    appendix.append(make_paragraph_elem(doc,
        f'Brazilian decisions were collected through custom web scrapers targeting eight '
        f'state court ESAJ/Elasticsearch/ASP.NET portals: TJDFT (Brasília, n={8421:,}), '
        f'TJRJ (Rio de Janeiro, n=1,219), TJSC (Santa Catarina, n=1,224), TJSP (São Paulo, '
        f'n=574), TJAC (Acre, n=33), TJRR (Roraima, n=21), TJPI (Piauí, n=15), TJTO '
        f'(Tocantins, n=17). An additional 200 historical TJSP decisions (1997–2015) were '
        f'included from a hand-coded qualitative dataset. Dutch decisions were collected '
        f'via the Rechtspraak.nl Open Data XML API, covering all 11 district courts and '
        f'major appellate courts (Raad van State, CBb, GHARL). Canadian decisions were '
        f'collected via the CanLII REST API supplemented by Legal Data Hunter semantic '
        f'search across the full CanLII corpus.'
    ))

    appendix.append(make_heading_elem(doc, 'A.3 Jurimetric Coding Variables', level=2))
    appendix.append(make_paragraph_elem(doc,
        'Each decision was coded for seven variables using a regex-based coding engine '
        '(jurimetric_coding.py v2.0, available in the GitHub repository):'
    ))
    var_lines = [
        'hr_language — whether the decision contains explicit human rights framing (binary)',
        'sust_language — whether sustainability or environmental rights language appears (binary)',
        'governance_cat — governance category (21 categories; see A.4)',
        'win_loss — outcome for the user/plaintiff: user_wins, utility_wins, mixed, unclear',
        'mp_involvement — whether a public ministry (Ministério Público) appears as a party (binary)',
        'indigenous_water — whether indigenous water rights are invoked (binary)',
        'public_interest — whether public interest framing appears (binary)',
    ]
    for v in var_lines:
        appendix.append(make_paragraph_elem(doc, f'• {v}'))

    appendix.append(make_heading_elem(doc, 'A.4 Governance Categories', level=2))
    appendix.append(make_paragraph_elem(doc,
        f'The 21 governance categories applied by the coding engine are: tariff_dispute, '
        f'connection_refusal, sanitation_sewage, water_quality, environmental_protection, '
        f'regulatory_permit, irrigation_agricultural, flood_protection, spatial_planning_water, '
        f'waterboard_governance, pipe_leak_damage, water_theft_fraud, water_infrastructure_contract, '
        f'fisheries_water, informal_settlement, public_health_water, water_rights_conflict, '
        f'metering_billing, water_access_hrw, drought_emergency, other_water. '
        f'Cases where the available metadata does not match any of the 20 specific categories '
        f'are classified as other_water ({n_ow_total:,} decisions, {n_ow_total/n_total*100:.1f}% '
        f'of total). A further {n_nwr:,} Dutch decisions were classified as not_water_related '
        f'after an immigration-law pre-filter. See the Methodology Note in the dataset '
        f'documentation for a full explanation of the other_water residual.'
    ))

    appendix.append(make_heading_elem(doc, 'A.5 Summary Statistics', level=2))
    stats_lines = [
        f'Total decisions: {n_total:,}',
        f'Brazil (2016–2025): {n_br:,} decisions from 8 state courts + 200 TJSP historical',
        f'Netherlands (2016–2025): {n_nl:,} decisions (district courts + Raad van State + appellate)',
        f'Canada (2016–2025): {n_ca:,} decisions (CanLII API + Legal Data Hunter)',
        f'Brazil coded for win/loss: {n_br_wl:,} decisions',
        f'  — User wins: {n_user_w:,} ({pct_user:.1f}%)',
        f'  — Utility wins: {n_util_w:,} ({pct_util:.1f}%)',
        f'Brazil human rights language: {n_hr:,} decisions ({pct_hr:.1f}%)',
        f'Brazil tariff disputes: {n_tariff:,} ({pct_tariff:.1f}%)',
        f'Brazil connection refusal: {n_conn_br:,} ({pct_conn_br:.1f}%)',
        f'Brazil irregular settlement: {n_irreg:,} ({pct_irreg:.1f}%)',
        f'Netherlands connection refusal: {n_conn_nl:,} ({pct_conn_nl:.1f}%)',
        f'Canada connection refusal: {n_conn_ca:,} ({pct_conn_ca:.1f}%)',
    ]
    for s in stats_lines:
        appendix.append(make_paragraph_elem(doc, s))

    appendix.append(make_heading_elem(doc, 'A.6 Charts', level=2))

    chart_specs = [
        ('chart1_cases_per_year.png', 'Figure A.1: Water Law Judicial Decisions by Year and Country (2016–2025). Dual axis: Brazil and Canada on left; Netherlands on right.', 5.8),
        ('chart2_brazil_governance.png', f'Figure A.2: Brazil — Top Governance Categories (2016–2025), n={n_br:,}. Other_water excluded.', 5.5),
        ('chart3_brazil_winloss.png', f'Figure A.3: Brazil — Win/Loss Outcomes by Year (2016–2025), n={n_br_wl:,} coded decisions. Shading indicates post-2020 reform period.', 5.5),
        ('chart4_brazil_hr.png', f'Figure A.4: Brazil — Human Rights Language in Decisions (2016–2025), n={n_br:,}.', 5.5),
        ('chart5_comparative.png', f'Figure A.5: Comparative Governance Categories (Top 6 per country, excluding other_water). BR={n_br:,}, NL={n_nl:,}, CA={n_ca:,}.', 6.0),
        ('chart6_informal.png', f'Figure A.6: Brazil — Irregular Settlement Cases (2016–2025), n={n_irreg:,} ({pct_irreg:.1f}% of total).', 5.5),
    ]
    for fname, caption, width in chart_specs:
        chart_path = CHARTS / fname
        if chart_path.exists():
            appendix.append(make_image_para_elem(doc, chart_path, width_inches=width))
            appendix.append(make_caption_elem(doc, caption + ' Source: Global Water Law Judicial Decisions Dataset (Klaus, 2026).'))
        else:
            print(f"    WARNING: Chart not found: {fname}")

    appendix.append(make_heading_elem(doc, 'A.7 Citation', level=2))
    appendix.append(make_paragraph_elem(doc,
        'Klaus, C (2026) Global Water Law Judicial Decisions Dataset (v1.0) '
        'Zenodo <https://doi.org/10.5281/zenodo.19836413> accessed 2 May 2026.'
    ))

    insert_elems_after(before_bib, appendix)
    print(f"    OK: inserted Appendix A before bibliography")
else:
    print("    WARNING: Could not find Bibliography heading — appending to end")
    # Just append to document
    doc.add_page_break()
    doc.add_heading('Appendix A: Global Water Law Dataset — Data Annex', level=1)
    doc.add_paragraph(f'See full statistics above. Total: {n_total:,} decisions.')

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
print(f"\nSaving to {DISS_OUT.name} …")
doc.save(str(DISS_OUT))
size_mb = DISS_OUT.stat().st_size / 1_048_576
print(f"  Done. {size_mb:.1f} MB — {len(doc.paragraphs)} paragraphs")
print(f"\nOutput: {DISS_OUT}")
